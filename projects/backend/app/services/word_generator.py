"""Word document generator for patent specifications."""

from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.api.patent_schemas import PatentSpecification


def generate_word_document(specification: PatentSpecification) -> BytesIO:
    """
    Generate a Word document from patent specification.

    Args:
        specification: Patent specification data

    Returns:
        BytesIO object containing the Word document
    """
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "맑은 고딕"
    font.size = Pt(11)

    # Title
    title = doc.add_heading(specification.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add sections
    _add_section(doc, "기술분야", specification.technical_field)
    _add_section(doc, "발명의 배경이 되는 기술", specification.background_art)
    _add_section(doc, "해결하려는 과제", specification.problem_to_solve)
    _add_section(doc, "과제의 해결 수단", specification.solution)
    _add_section(doc, "발명의 효과", specification.advantageous_effects)
    _add_section(
        doc, "발명을 실시하기 위한 구체적인 내용", specification.detailed_description
    )

    # Claims section
    doc.add_heading("청구항", level=1)
    for claim in specification.claims:
        claim_text = f"[청구항 {claim.number}]"
        if claim.depends_on:
            claim_text += f" (청구항 {claim.depends_on}에 종속)"
        claim_text += f"\n{claim.text}"

        p = doc.add_paragraph()
        p.add_run(f"청구항 {claim.number}").bold = True
        if claim.depends_on:
            p.add_run(f" (청구항 {claim.depends_on}에 종속)")
        p.add_run("\n" + claim.text)

    # Abstract
    _add_section(doc, "요약서", specification.abstract)

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


def _add_section(doc: Document, title: str, content: str) -> None:
    """
    Add a section with title and content to the document.

    Args:
        doc: Document object
        title: Section title
        content: Section content
    """
    doc.add_heading(title, level=1)
    # Split content by newlines and add as paragraphs
    for paragraph_text in content.split("\n"):
        if paragraph_text.strip():
            doc.add_paragraph(paragraph_text.strip())
